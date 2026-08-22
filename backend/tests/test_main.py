import asyncio
import json
from pathlib import Path

from docx import Document
from fastapi import HTTPException

import backend.main as backend_main


def paragraph_texts(docx_path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(docx_path).paragraphs]


def write_template(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


class DummyUpload:
    def __init__(self, filename: str, content: str):
        self.filename = filename
        self._content = content.encode("utf-8")

    async def read(self) -> bytes:
        return self._content


def test_health_endpoint_returns_ok():
    assert backend_main.health() == {"status": "ok"}


def test_list_templates_exposes_known_templates():
    templates = backend_main.list_templates()
    names = {item["id"] for item in templates}
    assert {"reference", "cloudbility-long", "cloudbility-short", "yuanchuangli-long", "yuanchuangli-short"} <= names
    assert any(item["supports_subtitle"] is True for item in templates if item["variant"] == "short")
    assert all(item["preview"].startswith("/template-covers/") for item in templates)


async def read_stream_events(response) -> list[dict]:
    chunks: list[str] = []
    async for chunk in response.body_iterator:
        chunks.append(chunk.decode("utf-8") if isinstance(chunk, bytes) else chunk)
    return [json.loads(line) for line in "".join(chunks).splitlines() if line.strip()]


def test_formalize_endpoint_returns_cleaned_markdown_defaults():
    upload = DummyUpload("demo.md", "# 系统说明书\n\n## 概述\n\n正文\n")

    response = asyncio.run(backend_main.formalize_markdown_endpoint(markdown_file=upload, file=None))
    events = asyncio.run(read_stream_events(response))
    payload = events[-1]["result"]

    assert payload["title"] == "系统说明书"
    assert payload["header_title"] == "系统说明书"
    assert payload["output_name"] == "系统说明书.docx"
    assert payload["cleaned_markdown"].startswith("# 系统说明书")
    assert payload["preview"] == payload["cleaned_markdown"]
    assert payload["file_name"] == "demo.md"
    assert [event["type"] for event in events[:-1]] == ["stage"] * 10


def test_convert_endpoint_uses_long_template(tmp_path: Path, monkeypatch):
    template_path = tmp_path / "long-template.docx"
    write_template(template_path, ["品牌 {{document_title}}", "{{main_content}}"])
    monkeypatch.setitem(backend_main.TEMPLATE_CHOICES, "test-long", template_path)

    upload = DummyUpload("input.md", "# 系统说明书\n\n## 概述\n\n正文\n")
    response = asyncio.run(
        backend_main.convert_markdown_endpoint(
            markdown_file=upload,
            file=None,
            template_id="test-long",
            title="",
            header_title="系统说明书",
            output_name="",
            subtitle="",
        )
    )

    out = tmp_path / "long-output.docx"
    out.write_bytes(response.body)
    texts = paragraph_texts(out)
    assert "品牌 系统说明书" in texts
    assert "概述" in texts
    assert "正文" in texts
    assert "{{main_content}}" not in texts


def test_convert_endpoint_uses_short_template(tmp_path: Path, monkeypatch):
    template_path = tmp_path / "short-template.docx"
    write_template(template_path, ["{{title}}", "{{subtitle}}", "{{main_content}}"])
    monkeypatch.setitem(backend_main.TEMPLATE_CHOICES, "test-short", template_path)

    upload = DummyUpload("input.md", "# 系统说明书\n\n## 概述\n\n正文\n")
    response = asyncio.run(
        backend_main.convert_markdown_endpoint(
            markdown_file=upload,
            file=None,
            template_id="test-short",
            title="",
            header_title="系统说明书",
            output_name="short-result.docx",
            subtitle="副标题",
        )
    )

    out = tmp_path / "short-output.docx"
    out.write_bytes(response.body)
    texts = paragraph_texts(out)
    assert "系统说明书" in texts
    assert "副标题" in texts
    assert "概述" in texts
    assert "正文" in texts
    assert "{{title}}" not in texts
    assert "{{subtitle}}" not in texts
    assert "{{main_content}}" not in texts
    assert 'filename="short-result.docx"' in response.headers["Content-Disposition"]


def test_convert_endpoint_rejects_unknown_template(tmp_path: Path):
    upload = DummyUpload("input.md", "# 系统说明书\n")

    try:
        asyncio.run(
            backend_main.convert_markdown_endpoint(
                markdown_file=upload,
                file=None,
                template_id="missing-template",
                title="",
                header_title="",
                output_name="",
                subtitle="",
            )
        )
    except HTTPException as exc:
        assert exc.status_code == 400
        assert exc.detail["error"] == "unknown_template"
    else:
        raise AssertionError("Expected HTTPException")


def test_convert_endpoint_ignores_subtitle_for_long_template(tmp_path: Path, monkeypatch):
    template_path = tmp_path / "long-template.docx"
    write_template(template_path, ["{{document_title}}", "{{subtitle}}", "{{main_content}}"])
    monkeypatch.setitem(backend_main.TEMPLATE_CHOICES, "test-long", template_path)

    upload = DummyUpload("input.md", "# 系统说明书\n\n## 概述\n\n正文\n")
    response = asyncio.run(
        backend_main.convert_markdown_endpoint(
            markdown_file=upload,
            file=None,
            template_id="test-long",
            title="系统说明书",
            header_title="系统说明书",
            output_name="",
            subtitle="不会生效",
        )
    )

    out = tmp_path / "long-output.docx"
    out.write_bytes(response.body)
    texts = paragraph_texts(out)
    assert "不会生效" not in texts
