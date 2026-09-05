from pathlib import Path
from zipfile import ZipFile

from docx import Document

from backend.md2word.converter import clean_image_caption, render_markdown_to_document
from backend.md2word.template_registry import get_style_map, get_template_path
from backend.md2word.workflow import convert_markdown_to_docx


def paragraph_texts(docx_path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(docx_path).paragraphs]


def paragraph_numbering(paragraph) -> tuple[int | None, int | None]:
    p_pr = paragraph._p.pPr
    num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
    num_id = getattr(num_pr, "numId", None) if num_pr is not None else None
    ilvl = getattr(num_pr, "ilvl", None) if num_pr is not None else None
    return (
        int(num_id.val) if num_id is not None and num_id.val is not None else None,
        int(ilvl.val) if ilvl is not None and ilvl.val is not None else None,
    )


def test_clean_image_caption_removes_redundant_prefixes():
    assert clean_image_caption("图1 系统架构", "images/arch.png") == "系统架构"
    assert clean_image_caption("image", "images/login-page.png") == "login-page"
    assert clean_image_caption("", "images/empty-caption.jpg") == "empty-caption"
    assert clean_image_caption("截图", "images/screen.png") == "截图"


def test_render_markdown_to_document_handles_headings_lists_and_paragraphs():
    doc = Document()

    render_markdown_to_document(
        doc,
        "# 系统说明书\n\n## 概述\n\n正文\n\n- 列表项\n\n1. 步骤一\n",
    )

    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "系统说明书" in texts
    assert "概述" in texts
    assert "正文" in texts
    assert "列表项" in texts
    assert "步骤一" in texts


def test_render_markdown_to_document_adds_missing_image_placeholder(tmp_path: Path):
    doc = Document()

    render_markdown_to_document(doc, "![截图](missing.png)", source_path=tmp_path / "doc.md")

    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "[图片缺失: missing.png]" in texts
    assert "图1 截图" in texts


def test_render_markdown_to_document_restarts_ordered_lists_per_block():
    doc = Document(str(get_template_path("reference")))

    render_markdown_to_document(
        doc,
        "## 第一节\n\n1. 步骤一\n2. 步骤二\n\n## 第二节\n\n1. 步骤三\n2. 步骤四\n",
        style_map=get_style_map("reference"),
    )

    numbered = {paragraph.text: paragraph_numbering(paragraph) for paragraph in doc.paragraphs if paragraph.text.startswith("步骤")}
    first_num_id, first_ilvl = numbered["步骤一"]
    second_num_id, second_ilvl = numbered["步骤三"]

    assert first_num_id is not None
    assert second_num_id is not None
    assert first_num_id != second_num_id
    assert first_ilvl == 0
    assert second_ilvl == 0


def test_render_markdown_to_document_keeps_reference_ordered_list_with_blank_lines_in_one_block():
    doc = Document(str(get_template_path("reference")))

    render_markdown_to_document(
        doc,
        "## 第一节\n\n1. 步骤一\n\n2. 步骤二\n\n3. 步骤三\n",
        style_map=get_style_map("reference"),
    )

    numbered = {paragraph.text: paragraph_numbering(paragraph) for paragraph in doc.paragraphs if paragraph.text.startswith("步骤")}
    first_num_id, first_ilvl = numbered["步骤一"]
    second_num_id, second_ilvl = numbered["步骤二"]
    third_num_id, third_ilvl = numbered["步骤三"]

    assert first_num_id is not None
    assert first_num_id == second_num_id == third_num_id
    assert first_ilvl == second_ilvl == third_ilvl == 0


def test_convert_markdown_to_docx_writes_docx_and_uses_baseline_formalizer(tmp_path: Path):
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    md_path.write_text(
        "# 系统说明书\n\n"
        "封面噪声\n\n"
        "## 目录\n\n"
        "目录内容\n\n"
        "## 第一章 概述\n\n"
        "正文 A\\.B\n",
        encoding="utf-8",
    )

    result = convert_markdown_to_docx(md_path, output_path)

    assert result.output_path == output_path
    assert output_path.exists()
    assert result.document_title == "系统说明书"
    texts = paragraph_texts(output_path)
    assert "系统说明书" in texts
    assert "概述" in texts
    assert "正文A.B" in texts
    assert "目录" in texts
    assert "封面噪声" in texts
    assert texts.count("系统说明书") == 1


def test_convert_markdown_to_docx_supports_ai_enhanced_mode(tmp_path: Path):
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    md_path.write_text(
        "# 系统说明书\n\n"
        "封面噪声\n\n"
        "## 目录\n\n"
        "目录内容\n\n"
        "## 第一章 概述\n\n"
        "正文 A\\.B\n",
        encoding="utf-8",
    )

    result = convert_markdown_to_docx(md_path, output_path, mode="ai_enhanced")

    assert result.output_path == output_path
    texts = paragraph_texts(output_path)
    assert "系统说明书" in texts
    assert "概述" in texts
    assert "正文A.B" in texts
    assert "目录" not in texts
    assert "封面噪声" not in texts


def test_convert_markdown_to_docx_embeds_existing_image(tmp_path: Path):
    md_path = tmp_path / "input.md"
    image_path = tmp_path / "pixel.png"
    output_path = tmp_path / "image.docx"
    image_path.write_bytes(
        bytes.fromhex(
            "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753de"
            "0000000c49444154789c63606060000000040001f61738550000000049454e44ae426082"
        )
    )
    md_path.write_text("# 图像文档\n\n## 图像\n\n![截图](pixel.png)\n", encoding="utf-8")

    convert_markdown_to_docx(md_path, output_path)

    with ZipFile(output_path) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media_files
    assert "图1 截图" in paragraph_texts(output_path)


def test_convert_markdown_to_docx_supports_long_template_placeholders(tmp_path: Path):
    template_path = tmp_path / "long-template.docx"
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    template = Document()
    template.add_paragraph("品牌 {{document_title}}")
    template.add_paragraph("{{main_content}}")
    template.save(template_path)
    md_path.write_text("# 系统说明书\n\n## 概述\n\n正文\n", encoding="utf-8")

    convert_markdown_to_docx(md_path, output_path, template_path=template_path)

    texts = paragraph_texts(output_path)
    assert "品牌 系统说明书" in texts
    assert "概述" in texts
    assert "正文" in texts
    assert "{{main_content}}" not in texts


def test_convert_markdown_to_docx_supports_short_template_placeholders(tmp_path: Path):
    template_path = tmp_path / "short-template.docx"
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    template = Document()
    template.add_paragraph("{{title}}")
    template.add_paragraph("{{subtitle}}")
    template.add_paragraph("{{main_content}}")
    template.save(template_path)
    md_path.write_text("# 系统说明书\n\n## 概述\n\n正文\n", encoding="utf-8")

    convert_markdown_to_docx(md_path, output_path, subtitle="副标题", template_path=template_path)

    texts = paragraph_texts(output_path)
    assert "系统说明书" in texts
    assert "副标题" in texts
    assert "概述" in texts
    assert "正文" in texts
    assert "{{title}}" not in texts
    assert "{{subtitle}}" not in texts
    assert "{{main_content}}" not in texts


def test_convert_markdown_to_docx_maps_heading_levels_for_template(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    template = Document()
    template.add_paragraph("{{document_title}}")
    template.add_paragraph("{{main_content}}")
    template.save(template_path)
    md_path.write_text("# 系统说明书\n\n## 一、项目背景\n\n### 2.1 原创方案\n", encoding="utf-8")

    convert_markdown_to_docx(md_path, output_path, template_path=template_path)

    texts = paragraph_texts(output_path)
    assert "系统说明书" in texts
    assert "项目背景" in texts
    assert "原创方案" in texts
    assert any(p.style.name == "Heading 1" and p.text == "项目背景" for p in Document(output_path).paragraphs)
    assert any(p.style.name == "Heading 2" and p.text == "原创方案" for p in Document(output_path).paragraphs)


def test_convert_markdown_to_docx_replaces_textbox_placeholder(tmp_path: Path):
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    md_path.write_text("# 系统说明书\n\n## 概述\n\n正文\n", encoding="utf-8")

    convert_markdown_to_docx(
        md_path,
        output_path,
        template_path=get_template_path("cloudbility-long"),
    )

    with ZipFile(output_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", "ignore")
    assert "{{document_title}}" not in xml
    assert "系统说明书" in xml
    assert any(p.style.name == "Cloudbility-正文" for p in Document(output_path).paragraphs if p.text.strip())


def test_convert_markdown_to_docx_uses_branded_template_style_map(tmp_path: Path):
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    md_path.write_text(
        "# 系统说明书\n\n"
        "## 概述\n\n"
        "正文\n\n"
        "- 列表项\n\n"
        "> 引用内容\n",
        encoding="utf-8",
    )

    convert_markdown_to_docx(
        md_path,
        output_path,
        template_path=get_template_path("cloudbility-long"),
    )

    paragraphs = [p for p in Document(output_path).paragraphs if p.text.strip()]
    styles_by_text = {p.text: p.style.name for p in paragraphs}
    assert styles_by_text["概述"] == "Heading 1"
    assert styles_by_text["正文"] == "Cloudbility-正文"
    assert styles_by_text["列表项"] == "Cloudbility-列表样式1级"
    assert styles_by_text["引用内容"] == "灰色文字"


def test_convert_markdown_to_docx_uses_reference_template_style_map(tmp_path: Path):
    md_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    md_path.write_text(
        "# 系统说明书\n\n"
        "## 概述\n\n"
        "正文\n\n"
        "> 引用内容\n\n"
        "```python\nprint('x')\n```\n\n"
        "| 列1 | 列2 |\n"
        "| --- | --- |\n"
        "| A | B |\n",
        encoding="utf-8",
    )

    convert_markdown_to_docx(
        md_path,
        output_path,
        template_path=get_template_path("reference"),
    )

    doc = Document(output_path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    styles_by_text = {p.text: p.style.name for p in paragraphs}
    assert styles_by_text["概述"] == "Heading 1"
    assert styles_by_text["正文"] == "Normal"
    assert styles_by_text["引用内容"] == "引用块"
    assert styles_by_text["print('x')"] == "代码块"
    assert doc.tables
    assert doc.tables[0].style.name == "CyanScript Table"
