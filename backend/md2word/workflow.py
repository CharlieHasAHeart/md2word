from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .converter import render_markdown_to_document
from .formalizer import ProcessingMode, formalize_markdown
from .template_registry import DEFAULT_STYLE_MAP, get_style_map_for_template_path
from .template_ops import (
    find_body_placeholder,
    remove_paragraphs_containing_token,
    replace_template_placeholders,
    replace_textbox_placeholder,
)


@dataclass(frozen=True)
class ConversionResult:
    output_path: Path
    document_title: str
    markdown_text: str


def convert_markdown_to_docx(
    md_path: str | Path,
    output_path: str | Path,
    document_title: str = "",
    subtitle: str = "",
    template_path: str | Path | None = None,
    mode: ProcessingMode = "baseline",
) -> ConversionResult:
    source = Path(md_path)
    target = Path(output_path)
    md_text = source.read_text(encoding="utf-8")
    formalized = formalize_markdown(md_text, source_path=source, mode=mode)
    title = document_title.strip() or formalized.document_title or source.stem
    body_markdown = _strip_document_title_heading(formalized.markdown_text)

    if template_path is None:
        doc = Document()
        heading = doc.add_paragraph(title, style=DEFAULT_STYLE_MAP.document_title)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        render_markdown_to_document(doc, body_markdown, source, style_map=DEFAULT_STYLE_MAP)
    else:
        doc = Document(str(template_path))
        style_map = get_style_map_for_template_path(template_path)
        replace_textbox_placeholder(doc, "{{document_title}}", title)
        if subtitle.strip():
            replace_textbox_placeholder(doc, "{{subtitle}}", subtitle.strip())
        replace_template_placeholders(
            doc,
            {
                "{{document_title}}": title,
                "{{title}}": title,
                "{{subtitle}}": subtitle.strip(),
            },
        )
        anchor = find_body_placeholder(doc, "{{main_content}}")
        if anchor is None:
            raise ValueError("Template placeholder {{main_content}} was not found.")
        render_markdown_to_document(
            doc,
            body_markdown,
            source,
            anchor=anchor,
            style_map=style_map,
        )
        remove_paragraphs_containing_token(doc, "{{main_content}}")

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return ConversionResult(output_path=target, document_title=title, markdown_text=formalized.markdown_text)


def _strip_document_title_heading(md_text: str) -> str:
    lines = md_text.splitlines()
    output: list[str] = []
    skipped = False
    for line in lines:
        if not skipped and re.match(r"^#\s+.+?\s*$", line):
            skipped = True
            continue
        output.append(line)
    return "\n".join(output).rstrip("\n") + ("\n" if output else "")
