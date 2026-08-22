from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.numbering import CT_Numbering
from docx.shared import Inches

from .formalizer import SUPPORTED_IMAGE_EXTENSIONS
from .template_registry import DEFAULT_STYLE_MAP, MarkdownStyleMap


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
IMAGE_RE = re.compile(r"^!\[(?P<caption>[^\]]*)\]\((?P<target>[^)]+)\)\s*$")
UNORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)[-+*]\s+(?P<text>.+?)\s*$")
ORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)\d+[.)]\s+(?P<text>.+?)\s*$")
BLOCKQUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
FENCE_RE = re.compile(r"^\s*(`{3,}|~{3,})(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def render_markdown_to_document(
    doc: Document,
    md_text: str,
    source_path: str | Path | None = None,
    anchor=None,
    style_map: MarkdownStyleMap | None = None,
) -> None:
    styles = style_map or DEFAULT_STYLE_MAP
    base_dir = Path(source_path).resolve().parent if source_path else Path.cwd()
    figure_index = 0
    lines = md_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        fence_match = FENCE_RE.match(line)
        if fence_match:
            code_block, index = _consume_fenced_block(lines, index, fence_match.group(1))
            _move_before_anchor(_add_paragraph(doc, code_block, styles.code_block or styles.body), anchor)
            continue

        if _is_markdown_table_block(lines, index):
            table_rows, index = _consume_markdown_table(lines, index)
            table = _add_table(doc, table_rows, styles.table)
            _move_table_before_anchor(table, anchor)
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            style_name = _style_for_heading_level(len(heading_match.group(1)), styles)
            _move_before_anchor(_add_paragraph(doc, heading_match.group(2).strip(), style_name), anchor)
            index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            figure_index += 1
            _add_image_or_placeholder(
                doc,
                base_dir,
                image_match.group("target"),
                image_match.group("caption"),
                figure_index,
                anchor=anchor,
                style_map=styles,
            )
            index += 1
            continue

        blockquote_match = BLOCKQUOTE_RE.match(line)
        if blockquote_match:
            quote_text, index = _consume_blockquote(lines, index)
            _move_before_anchor(_add_paragraph(doc, quote_text, styles.blockquote or styles.body), anchor)
            continue

        unordered_match = UNORDERED_LIST_RE.match(line)
        if unordered_match:
            list_style = styles.bullet_list_2 if _list_level(unordered_match.group("indent")) > 1 else styles.bullet_list
            _move_before_anchor(_add_paragraph(doc, unordered_match.group("text").strip(), list_style or styles.body), anchor)
            index += 1
            continue

        ordered_match = ORDERED_LIST_RE.match(line)
        if ordered_match:
            ordered_items, index = _consume_ordered_list_block(
                lines,
                index,
                allow_blank_lines=_uses_reference_ordered_list_handling(styles),
            )
            _render_ordered_list_block(doc, ordered_items, styles, anchor=anchor)
            continue

        _move_before_anchor(_add_paragraph(doc, stripped, styles.body), anchor)
        index += 1


def _add_image_or_placeholder(
    doc: Document,
    base_dir: Path,
    target: str,
    caption: str,
    figure_index: int,
    anchor=None,
    style_map: MarkdownStyleMap | None = None,
) -> None:
    styles = style_map or DEFAULT_STYLE_MAP
    clean_target = target.strip().strip("<>")
    image_path = Path(clean_target)
    if not image_path.is_absolute():
        image_path = base_dir / clean_target

    if image_path.suffix.lower() in SUPPORTED_IMAGE_EXTENSIONS and image_path.exists():
        paragraph = _add_paragraph(doc, "", styles.image)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
        _move_before_anchor(paragraph, anchor)
    else:
        _move_before_anchor(_add_paragraph(doc, f"[图片缺失: {clean_target}]", styles.body), anchor)

    caption_text = clean_image_caption(caption, image_path)
    caption_style = styles.image_caption or styles.body
    caption_paragraph = _add_paragraph(doc, f"图{figure_index} {caption_text}".rstrip(), caption_style)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _move_before_anchor(caption_paragraph, anchor)


def clean_image_caption(caption: str, image_path: str | Path) -> str:
    text = caption.strip()
    text = re.sub(r"^(?:图\s*\d+|图片|截图|Figure\s*\d+)\s*[:：.、-]?\s*", "", text, flags=re.I).strip()
    if not text and caption.strip() in {"图片", "截图"}:
        return caption.strip()
    if not text or text.lower() in {"image", "img", "figure", "screenshot", "photo", "picture"}:
        return Path(image_path).stem
    return text


def _move_before_anchor(paragraph, anchor) -> None:
    if anchor is not None:
        anchor.addprevious(paragraph._element)


def _move_table_before_anchor(table, anchor) -> None:
    if anchor is not None:
        anchor.addprevious(table._element)


def _add_paragraph(doc: Document, text: str, style: str | None):
    if not style:
        return doc.add_paragraph(text)
    try:
        return doc.add_paragraph(text, style=style)
    except KeyError:
        return doc.add_paragraph(text)


def _add_table(doc: Document, rows: list[list[str]], style: str | None):
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    if style:
        try:
            table.style = style
        except KeyError:
            pass
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    return table


def _style_for_heading_level(level: int, style_map: MarkdownStyleMap) -> str | None:
    if level <= 1:
        return style_map.document_title or style_map.heading_1 or style_map.body
    if level == 2:
        return style_map.heading_1 or style_map.body
    if level == 3:
        return style_map.heading_2 or style_map.body
    return style_map.heading_3 or style_map.body


def _list_level(indent: str) -> int:
    expanded = indent.replace("\t", "    ")
    return 2 if len(expanded) >= 2 else 1


def _consume_ordered_list_block(
    lines: list[str],
    start_index: int,
    allow_blank_lines: bool = False,
) -> tuple[list[tuple[int, str]], int]:
    items: list[tuple[int, str]] = []
    index = start_index
    root_level: int | None = None
    while index < len(lines):
        line = lines[index]
        match = ORDERED_LIST_RE.match(line)
        if not match:
            if allow_blank_lines and not line.strip():
                next_index = index + 1
                while next_index < len(lines) and not lines[next_index].strip():
                    next_index += 1
                if next_index < len(lines):
                    next_match = ORDERED_LIST_RE.match(lines[next_index])
                    if next_match is not None:
                        next_level = _list_level(next_match.group("indent"))
                        if root_level is None or next_level == root_level:
                            index = next_index
                            continue
            break
        level = _list_level(match.group("indent"))
        if root_level is None:
            root_level = level
        items.append((level, match.group("text").strip()))
        index += 1
    return items, index


def _render_ordered_list_block(
    doc: Document,
    items: list[tuple[int, str]],
    styles: MarkdownStyleMap,
    anchor=None,
) -> None:
    if not items:
        return

    levels = {max(level - 1, 0) for level, _ in items}
    restart_num_id = _create_restart_numbering(doc, styles.ordered_list, levels)

    for level, text in items:
        paragraph_style = _paragraph_style_for_ordered_item(styles, level)
        paragraph = _add_paragraph(doc, text, paragraph_style)
        if restart_num_id is not None:
            _apply_numbering_override(paragraph, restart_num_id, max(level - 1, 0))
        _move_before_anchor(paragraph, anchor)


def _create_restart_numbering(
    doc: Document,
    style_name: str | None,
    levels: set[int],
) -> int | None:
    base_num_id = _style_num_id(doc, style_name)
    if base_num_id is None:
        return None

    numbering_part = getattr(doc.part, "numbering_part", None)
    if numbering_part is None:
        return None

    numbering = numbering_part.element
    if not isinstance(numbering, CT_Numbering):
        return None

    try:
        base_num = numbering.num_having_numId(base_num_id)
    except KeyError:
        return None

    abstract_num_id = int(base_num.abstractNumId.val)
    new_num = numbering.add_num(abstract_num_id)
    for ilvl in sorted(levels):
        new_num.add_lvlOverride(ilvl).add_startOverride(1)
    return int(new_num.numId)


def _style_num_id(doc: Document, style_name: str | None) -> int | None:
    if not style_name:
        return None
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None

    p_pr = getattr(style.element, "pPr", None)
    num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
    num_id = getattr(num_pr, "numId", None) if num_pr is not None else None
    if num_id is None or num_id.val is None:
        return None
    return int(num_id.val)


def _apply_numbering_override(paragraph, num_id: int, ilvl: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_numId().val = num_id
    num_pr.get_or_add_ilvl().val = ilvl


def _paragraph_style_for_ordered_item(styles: MarkdownStyleMap, level: int) -> str | None:
    # For the reference template, the numbered-list paragraph style already binds its
    # own numId. Reusing that style while also overriding numId at paragraph level
    # causes Word to render every item as "1". Keep Normal text styling there and
    # let numbering.xml drive the visible list formatting.
    if _uses_reference_ordered_list_handling(styles):
        return styles.body
    return styles.ordered_list_2 if level > 1 else styles.ordered_list


def _uses_reference_ordered_list_handling(styles: MarkdownStyleMap) -> bool:
    return styles.ordered_list == "列表-有序"


def _consume_fenced_block(lines: list[str], start_index: int, marker: str) -> tuple[str, int]:
    output: list[str] = []
    opener = marker[0]
    index = start_index + 1
    while index < len(lines):
        line = lines[index]
        if re.match(rf"^\s*{re.escape(opener)}{{3,}}\s*$", line):
            return "\n".join(output), index + 1
        output.append(line)
        index += 1
    return "\n".join(output), index


def _consume_blockquote(lines: list[str], start_index: int) -> tuple[str, int]:
    output: list[str] = []
    index = start_index
    while index < len(lines):
        match = BLOCKQUOTE_RE.match(lines[index])
        if not match:
            break
        output.append(match.group(1).strip())
        index += 1
    return "\n".join(output).strip(), index


def _is_markdown_table_block(lines: list[str], start_index: int) -> bool:
    if start_index + 1 >= len(lines):
        return False
    return _looks_like_table_row(lines[start_index]) and TABLE_SEPARATOR_RE.match(lines[start_index + 1]) is not None


def _consume_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows = [_parse_table_row(lines[start_index])]
    index = start_index + 2
    while index < len(lines) and _looks_like_table_row(lines[index]):
        rows.append(_parse_table_row(lines[index]))
        index += 1
    return rows, index


def _looks_like_table_row(line: str) -> bool:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return False
    return not any(
        pattern.match(line)
        for pattern in (HEADING_RE, IMAGE_RE, UNORDERED_LIST_RE, ORDERED_LIST_RE, BLOCKQUOTE_RE, FENCE_RE)
    )


def _parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


__all__ = [
    "clean_image_caption",
    "render_markdown_to_document",
]
